import json
import re
import uuid
from io import BytesIO
import anthropic
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from pyvis.network import Network
from snowflake.snowpark import Session
import networkx as nx
from docx import Document
from docx.shared import Pt, RGBColor, Inches


st.markdown("""
<style>
/* hide sidebar page nav (top buttons handle navigation) */
[data-testid="stSidebarNav"] { display:none !important; }

/* ----- sidebar: NUS-blue scheme (matches Industry Collaboration Overview) ----- */
[data-testid="stSidebar"]            { background:#003D7C !important; }
[data-testid="stSidebar"] *          { color:#D6E4F0 !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] label      { color:#FFFFFF !important; }

/* inputs, selects, textareas */
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea,
[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background:#0A3060 !important; color:#FFFFFF !important;
    border:1px solid #1A5090 !important; border-radius:6px !important;
}

/* buttons */
[data-testid="stSidebar"] button,
[data-testid="stSidebar"] .stButton > button {
    background:#2E6DA4 !important; color:#FFFFFF !important; border:none !important;
    border-radius:6px !important; font-weight:600 !important; font-size:14px !important;
}
[data-testid="stSidebar"] button:hover,
[data-testid="stSidebar"] .stButton > button:hover {
    background:#245888 !important; color:#FFFFFF !important;
}
</style>
""", unsafe_allow_html=True)

c1, c2, _ = st.columns([2, 2, 6])
with c1:
    st.button("🧲 Potential Collaborator Finder", disabled=True, use_container_width=True)
with c2:
    if st.button("📊 Industry Collaboration Overview", use_container_width=True):
        st.switch_page("pages/2_Industry_Intelligence_Dashboard.py")

st.title("🧲 Potential Collaborator Finder")
st.write("Discover institutions and corporations with shared research interests for potential collaboration.")


# -----------------------------
# Snowflake connection
# -----------------------------
def _load_private_key(pem_str: str) -> bytes:
    """Load a PEM private key string and return DER bytes for Snowflake auth."""
    from cryptography.hazmat.primitives import serialization
    p_key = serialization.load_pem_private_key(pem_str.encode(), password=None)
    return p_key.private_bytes(
        encoding=serialization.Encoding.DER,
        format=serialization.PrivateFormat.PKCS8,
        encryption_algorithm=serialization.NoEncryption(),
    )


@st.cache_resource
def create_session():
    connection_parameters = {
        "account": st.secrets["snowflake"]["account"],
        "user": st.secrets["snowflake"]["user"],
        "private_key": _load_private_key(st.secrets["snowflake"]["private_key"]),
        "role": st.secrets["snowflake"]["role"],
        "warehouse": st.secrets["snowflake"]["warehouse"],
        "database": st.secrets["snowflake"]["database"],
        "schema": st.secrets["snowflake"]["schema"],
    }
    return Session.builder.configs(connection_parameters).create()


session = create_session()

client = anthropic.Anthropic(api_key=st.secrets["anthropic"]["api_key"])

# -----------------------------
# Helper functions
# -----------------------------
@st.cache_data(ttl=3600)
def run_query(sql: str) -> pd.DataFrame:
    return session.sql(sql).to_pandas()


def sql_escape(value: str) -> str:
    return value.replace("'", "''")


def _log_session_id() -> str:
    if "log_session_id" not in st.session_state:
        st.session_state.log_session_id = uuid.uuid4().hex[:12]
    return st.session_state.log_session_id


def log_query(page, user_query, response_type=None, subject_filter=None,
              existing_only=None, llm_answer=None, result_count=None, result_orgs=None):
    """Best-effort logging of a query + LLM output to QUERY_LOG. Never raises — a
    logging failure must not break the user's request."""
    try:
        def esc(v):
            if v is None:
                return "NULL"
            if isinstance(v, bool):
                return "TRUE" if v else "FALSE"
            if isinstance(v, (int, float)):
                return str(v)
            s = str(v).replace("\\", "\\\\").replace("'", "''")
            return "'" + s + "'"
        sql = (
            "INSERT INTO INDUSTRY_AGG.PUBLIC.QUERY_LOG "
            "(PAGE, USER_QUERY, RESPONSE_TYPE, SUBJECT_FILTER, EXISTING_ONLY, "
            "LLM_ANSWER, RESULT_COUNT, RESULT_ORGS, SESSION_ID) "
            f"SELECT {esc(page)}, {esc(user_query)}, {esc(response_type)}, {esc(subject_filter)}, "
            f"{esc(existing_only)}, {esc(llm_answer)}, {esc(result_count)}, {esc(result_orgs)}, "
            f"{esc(_log_session_id())}"
        )
        session.sql(sql).collect()
    except Exception:
        pass


def get_node_color(node_type: str, node_name: str, nus_affiliated) -> str:
    node_type = str(node_type).upper()
    node_name = str(node_name).upper()

    if str(nus_affiliated).upper() in ["TRUE", "1", "YES"]:
        return "#ff9933"

    if "NATIONAL UNIVERSITY OF SINGAPORE" in node_name:
        return "#ff9933"

    if "SUBJECT" in node_type:
        return "#F4B183"

    if "APPLICANT" in node_type:
        return "#ccccff"

    if "INSTITUTE" in node_type:
        return "#33cccc"

    return "#D9D9D9"


def keep_top_communities(df: pd.DataFrame, top_n: int = 10) -> pd.DataFrame:
    if df.empty:
        return df

    G = nx.Graph()

    for _, row in df.iterrows():
        G.add_edge(row["SOURCE"], row["TARGET"], weight=float(row["WEIGHT"]))

    if G.number_of_edges() == 0:
        return df

    communities = list(nx.community.greedy_modularity_communities(G, weight="weight"))

    community_rows = []
    for i, community in enumerate(communities):
        community_rows.append({
            "community_id": i,
            "nodes": set(community),
            "size": len(community),
        })

    community_rows = sorted(community_rows, key=lambda x: x["size"], reverse=True)
    top_communities = community_rows[:top_n]

    node_to_community = {}
    for community in top_communities:
        for node in community["nodes"]:
            node_to_community[node] = community["community_id"]

    top_nodes = set(node_to_community.keys())

    filtered_df = df[
        df["SOURCE"].isin(top_nodes) & df["TARGET"].isin(top_nodes)
    ].copy()
    filtered_df["CLUSTER"] = filtered_df["SOURCE"].map(node_to_community)

    return filtered_df


def keep_top_n_neighbours(df: pd.DataFrame, search_term: str, top_n: int) -> pd.DataFrame:
    """Keep only the top N neighbours by total edge weight connected to the search term node."""
    if df.empty or not search_term:
        return df

    term = search_term.upper()

    mask = (
        df["SOURCE_NAME"].str.upper().str.contains(term, regex=False) |
        df["TARGET_NAME"].str.upper().str.contains(term, regex=False)
    )
    direct_edges = df[mask].copy()

    if direct_edges.empty:
        return df

    def get_neighbour(row):
        if term in str(row["SOURCE_NAME"]).upper():
            return row["TARGET_NAME"]
        return row["SOURCE_NAME"]

    direct_edges["NEIGHBOUR"] = direct_edges.apply(get_neighbour, axis=1)

    top_neighbours = (
        direct_edges.groupby("NEIGHBOUR")["WEIGHT"]
        .sum()
        .sort_values(ascending=False)
        .head(top_n)
        .index.tolist()
    )

    return df[
        df["SOURCE_NAME"].isin(top_neighbours) |
        df["TARGET_NAME"].isin(top_neighbours)
    ].copy()


def run_similar_no_collab_query(
    institution: str,
    ip_type: str = None,
    category: str = None,
    top_n: int = 20,
    subject_filter: str = None,
) -> pd.DataFrame:
    """
    Find organisations with similar subject interests to the given institution
    that have NOT directly appeared in any edge with it.
    Returns a ranked table with ORG_NAME, ORG_CATEGORY, SHARED_SUBJECTS, TOTAL_WEIGHT.

    Actual edge types in data:
      Applicant_Subject    (patents  — APP:: nodes to SUBJ:: nodes)
      Institute_Subject    (publications — INST:: nodes to SUBJ:: nodes)
      Applicant_Applicant  (patent co-applicants)
      Institution_Institution (publication co-authors)
    IP_TYPE values: 'Patents', 'Publications'
    """
    safe_inst = sql_escape(institution)

    # IP filter — match actual values "Patents" / "Publications"
    ip_filter = f"AND IP_TYPE = '{sql_escape(ip_type)}'" if ip_type and ip_type != "All" else ""

    # Subject edge types depend on ip_type
    if ip_type == "Patents":
        subject_edge_types = "('Applicant_Subject')"
    elif ip_type == "Publications":
        subject_edge_types = "('Institute_Subject')"
    else:
        subject_edge_types = "('Applicant_Subject', 'Institute_Subject')"

    # Direct collab edge types depend on ip_type
    if ip_type == "Patents":
        collab_edge_types = "('Applicant_Applicant')"
    elif ip_type == "Publications":
        collab_edge_types = "('Institution_Institution')"
    else:
        collab_edge_types = "('Applicant_Applicant', 'Institution_Institution')"

    cat_filter = (
        f"AND (SOURCE_CATEGORY = '{sql_escape(category)}' OR TARGET_CATEGORY = '{sql_escape(category)}')"
        if category and category != "All" else ""
    )
    subject_clause = (
        f"AND (SOURCE_NAME ILIKE '%{sql_escape(subject_filter)}%' OR TARGET_NAME ILIKE '%{sql_escape(subject_filter)}%')"
        if subject_filter else ""
    )

    sql = f"""
WITH inst_subjects AS (
    -- Step 1: all subjects the institution is connected to
    SELECT DISTINCT
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET ELSE SOURCE END AS SUBJECT_ID
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE (SOURCE_NAME ILIKE '%{safe_inst}%' OR TARGET_NAME ILIKE '%{safe_inst}%')
    AND EDGE_TYPE IN {subject_edge_types}
    {ip_filter}
    {subject_clause}
),
org_subject_edges AS (
    -- Step 2: all orgs connected to those same subjects, excluding the institution itself
    SELECT
        CASE WHEN SOURCE_TYPE IN ('Applicant', 'Institutes') THEN SOURCE ELSE TARGET END AS ORG_ID,
        CASE WHEN SOURCE_TYPE IN ('Applicant', 'Institutes') THEN SOURCE_NAME ELSE TARGET_NAME END AS ORG_NAME,
        CASE WHEN SOURCE_TYPE IN ('Applicant', 'Institutes') THEN SOURCE_CATEGORY ELSE TARGET_CATEGORY END AS ORG_CATEGORY,
        CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN SOURCE
             ELSE TARGET END AS MATCHED_SUBJECT_ID,
        WEIGHT
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE EDGE_TYPE IN {subject_edge_types}
    {ip_filter}
    AND (SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) OR TARGET IN (SELECT SUBJECT_ID FROM inst_subjects))
    AND SOURCE_NAME NOT ILIKE '%{safe_inst}%'
    AND TARGET_NAME NOT ILIKE '%{safe_inst}%'
    {cat_filter}
),
org_matches AS (
    -- Step 3: aggregate — count distinct shared subjects and total weight per org
    SELECT
        ORG_ID,
        ORG_NAME,
        ORG_CATEGORY,
        COUNT(DISTINCT MATCHED_SUBJECT_ID) AS SHARED_SUBJECTS,
        SUM(WEIGHT) AS TOTAL_WEIGHT
    FROM org_subject_edges
    GROUP BY ORG_ID, ORG_NAME, ORG_CATEGORY
),
direct_collabs AS (
    -- Step 4: any org that has ever appeared in a direct collaboration edge with the institution
    SELECT DISTINCT
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET ELSE SOURCE END AS COLLAB_ID
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE (SOURCE_NAME ILIKE '%{safe_inst}%' OR TARGET_NAME ILIKE '%{safe_inst}%')
    AND EDGE_TYPE IN {collab_edge_types}
)
-- Step 5: return orgs with shared subjects that never directly collaborated
SELECT
    ORG_ID,
    ORG_NAME,
    ORG_CATEGORY,
    SHARED_SUBJECTS,
    TOTAL_WEIGHT
FROM org_matches
WHERE ORG_ID NOT IN (SELECT COLLAB_ID FROM direct_collabs)
AND SHARED_SUBJECTS > 0
ORDER BY SHARED_SUBJECTS DESC, TOTAL_WEIGHT DESC
LIMIT {top_n}
"""
    return run_query(sql)


# Map common informal terms to actual QS granular subjects (fallback when the LLM
# passes an informal term straight through)
_FIELD_ALIASES = {
    "AI": "COMPUTER SCIENCE",
    "ARTIFICIAL INTELLIGENCE": "COMPUTER SCIENCE",
    "MACHINE LEARNING": "COMPUTER SCIENCE",
    "ML": "COMPUTER SCIENCE",
    "DEEP LEARNING": "COMPUTER SCIENCE",
    "COMPUTER SCIENCE": "COMPUTER SCIENCE",
    "DATA SCIENCE": "DATA SCIENCE",
    "SEMICONDUCTORS": "ENGINEERING - ELECTRICAL",
    "SEMICONDUCTOR": "ENGINEERING - ELECTRICAL",
    "ELECTRONICS": "ENGINEERING - ELECTRICAL",
    "FINTECH": "COMPUTER SCIENCE | ACCOUNTING & FINANCE",
    "PAYMENTS": "COMPUTER SCIENCE | ACCOUNTING & FINANCE",
}


def _field_clause(subject: str, col: str = "p.QS_SUBJECT") -> str:
    """Build a SQL clause matching one or more QS subjects (separated by '|').
    Supports multi-subject mapping for cross-disciplinary topics."""
    raw = (subject or "").strip()
    if not raw:
        return "TRUE"
    terms = []
    for p in raw.split("|"):
        p = p.strip()
        if not p:
            continue
        mapped = _FIELD_ALIASES.get(p.upper(), p.upper())
        # an alias may itself expand to several subjects (e.g. fintech → CS | Finance)
        for m in mapped.split("|"):
            m = m.strip()
            if m and m not in terms:
                terms.append(m)
    if not terms:
        return "TRUE"
    ors = " OR ".join(f"CONTAINS(UPPER({col}), '{t.replace(chr(39), chr(39)*2)}')" for t in terms)
    return f"({ors})"


@st.cache_data(ttl=3600, show_spinner=False)
def run_recommendation_flat(
    subject: str = None,
    existing_only: bool = False,
    top_n: int = 10,
) -> pd.DataFrame:
    """Flat-table recommendation scoring on PAT_PUB + ENTITIES (NUS-centric).

    Subject-conditioned, category-weighted composite score:
      score = field_ip * (1+focus) * (0.5+0.5*recency) * category_weight * institute_bonus
    Category weights: Corporation/Government-Nonprofit 1.0, Hospital 0.8, Institute 0.6, Individual excluded.
    """
    field_clause = _field_clause(subject)
    # existing-only ("collaborate most") → literal count of joint works in the field;
    # otherwise ("recommend") → composite potential score over all orgs.
    order_col = "FCOLLAB" if existing_only else "NEW_SCORE"
    row_filter = ("IS_NEW_OPPORTUNITY = FALSE AND FCOLLAB > 0"
                  if existing_only else "TRUE")

    sql = f"""
WITH name_map AS (
    SELECT DISTINCT TRIM(UPPER(n.VALUE::STRING)) AS NNAME,
        PARENT_BRAND, CATEGORY, NUS_AFFILIATED
    FROM INDUSTRY_AGG.PUBLIC.ENTITIES,
         LATERAL FLATTEN(INPUT => SPLIT(
            COALESCE(NORMALIZED_NAMES_PAT,'') || '|' || COALESCE(NORMALIZED_NAMES_PUB,''), '|')) n
    WHERE TRIM(n.VALUE::STRING) <> '' AND CATEGORY <> 'Individual'
),
flat AS (
    SELECT DISTINCT p.UID, m.PARENT_BRAND AS ORG, m.CATEGORY AS CAT,
        m.NUS_AFFILIATED AS IS_NUS,
        {field_clause} AS IN_FIELD,
        p.APPLICATION_PUBLICATION_YEAR AS YR
    FROM INDUSTRY_AGG.PUBLIC.PAT_PUB p,
         LATERAL FLATTEN(INPUT => SPLIT(p.NORMALIZED_NAMES_CONCAT, '|')) f
    JOIN name_map m ON TRIM(UPPER(f.VALUE::STRING)) = m.NNAME
),
nus_uids AS (SELECT DISTINCT UID FROM flat WHERE IS_NUS),
yrange AS (SELECT MIN(YR) AS MINY, MAX(YR) AS MAXY FROM flat),
org_inst AS (
    -- academic-collaboration track record: distinct NON-NUS institutes the org has co-appeared with
    SELECT a.ORG, COUNT(DISTINCT b.ORG) AS N_INST
    FROM flat a JOIN flat b ON a.UID = b.UID AND a.ORG <> b.ORG
        AND b.CAT = 'Institute' AND b.IS_NUS = FALSE
    GROUP BY a.ORG
),
agg AS (
    SELECT f.ORG, ANY_VALUE(f.CAT) AS CAT, BOOLOR_AGG(f.IS_NUS) AS IS_NUS,
        COUNT(DISTINCT f.UID) AS TOTAL,
        COUNT(DISTINCT IFF(f.IN_FIELD, f.UID, NULL)) AS FIELD_CNT,
        COUNT(DISTINCT IFF(f.IN_FIELD AND f.UID IN (SELECT UID FROM nus_uids), f.UID, NULL)) AS FCOLLAB,
        MAX(IFF(f.UID IN (SELECT UID FROM nus_uids), 1, 0)) AS EXISTING_ANY,
        AVG(IFF(f.IN_FIELD,
            (f.YR - (SELECT MINY FROM yrange)) / NULLIF((SELECT MAXY FROM yrange) - (SELECT MINY FROM yrange), 0),
            NULL)) AS RECENCY_W
    FROM flat f GROUP BY f.ORG
),
scored AS (
    SELECT a.ORG AS ORG_NAME, a.CAT AS ORG_CATEGORY,
        NOT (a.EXISTING_ANY = 1) AS IS_NEW_OPPORTUNITY,
        a.FIELD_CNT, a.TOTAL, a.FCOLLAB, COALESCE(i.N_INST, 0) AS N_INST,
        a.FIELD_CNT / NULLIF(a.TOTAL,0) AS FOCUS,
        COALESCE(a.RECENCY_W, 0) AS RECENT,
        CASE a.CAT WHEN 'Corporation' THEN 1.0 WHEN 'Government / Non-profit' THEN 1.0
                   WHEN 'Hospital' THEN 0.8 WHEN 'Institute' THEN 0.6 ELSE 0 END AS CAT_W,
        1 + LEAST(0.5, 0.3 * LN(1 + COALESCE(i.N_INST,0))) AS INST_BONUS
    FROM agg a LEFT JOIN org_inst i ON a.ORG = i.ORG
    WHERE a.IS_NUS = FALSE AND a.FIELD_CNT > 0
      AND a.CAT IN ('Corporation','Government / Non-profit','Hospital','Institute')
)
SELECT ORG_NAME, ORG_CATEGORY, IS_NEW_OPPORTUNITY, FIELD_CNT, TOTAL, FCOLLAB, N_INST,
       ROUND(FOCUS,2) AS FOCUS, ROUND(RECENT,2) AS RECENT,
       ROUND(LN(1+FIELD_CNT) * (1+FOCUS) * (0.5+0.5*RECENT) * CAT_W * INST_BONUS, 2) AS NEW_SCORE,
       ROUND(LN(1+FCOLLAB)  * (1+FOCUS) * (0.5+0.5*RECENT) * CAT_W * INST_BONUS, 2) AS EXI_SCORE
FROM scored
WHERE {row_filter}
ORDER BY {order_col} DESC
LIMIT {int(top_n)}
"""
    df = run_query(sql)
    if not df.empty:
        _sc = "FCOLLAB" if existing_only else "NEW_SCORE"
        _mx = df[_sc].max()
        # normalise to 0–100 relative to the strongest match in this result set
        df["MATCH_SCORE"] = ((100 * df[_sc] / _mx).round().astype(int) if _mx else 0)
    return df


@st.cache_data(ttl=3600, show_spinner=False)
def run_titles_for_flat_orgs(org_names: tuple, subject: str = None) -> pd.DataFrame:
    """Fetch up to 3 recent sample titles per org (parent brand) per IP type, in the field."""
    if not org_names:
        return pd.DataFrame()
    field_clause = _field_clause(subject)
    orgs_esc = ",".join("'" + str(o).replace("'", "''") + "'" for o in org_names)
    sql = f"""
WITH name_map AS (
    SELECT DISTINCT TRIM(UPPER(n.VALUE::STRING)) AS NNAME, PARENT_BRAND
    FROM INDUSTRY_AGG.PUBLIC.ENTITIES,
         LATERAL FLATTEN(INPUT => SPLIT(
            COALESCE(NORMALIZED_NAMES_PAT,'') || '|' || COALESCE(NORMALIZED_NAMES_PUB,''), '|')) n
    WHERE PARENT_BRAND IN ({orgs_esc}) AND TRIM(n.VALUE::STRING) <> ''
),
matched AS (
    SELECT m.PARENT_BRAND AS ORG_NAME, p.IP_TYPE, p.TITLE,
        ROW_NUMBER() OVER (PARTITION BY m.PARENT_BRAND, p.IP_TYPE
                           ORDER BY p.APPLICATION_PUBLICATION_YEAR DESC) AS rn
    FROM INDUSTRY_AGG.PUBLIC.PAT_PUB p,
         LATERAL FLATTEN(INPUT => SPLIT(p.NORMALIZED_NAMES_CONCAT, '|')) f
    JOIN name_map m ON TRIM(UPPER(f.VALUE::STRING)) = m.NNAME
    WHERE {field_clause} AND p.TITLE IS NOT NULL AND p.TITLE <> ''
)
SELECT ORG_NAME, IP_TYPE, TITLE FROM matched WHERE rn <= 3 ORDER BY ORG_NAME, IP_TYPE, rn
"""
    return run_query(sql)


def generate_recommendations_flat(recs_df, subject=None, existing_only=False, titles_df=None):
    """Write recommendations from the flat-model results, framed around the new signals."""
    titles_by_org = {}
    if titles_df is not None and not titles_df.empty:
        for _, row in titles_df.iterrows():
            oid = str(row["ORG_NAME"]); ip = str(row["IP_TYPE"])
            t = str(row["TITLE"]).strip().title()
            titles_by_org.setdefault(oid, {"Patents": [], "Publications": []})
            if ip in titles_by_org[oid]:
                titles_by_org[oid][ip].append(t)

    rows = []
    for _, r in recs_df.iterrows():
        org = str(r["ORG_NAME"])
        tier = "🆕 New Opportunity" if r["IS_NEW_OPPORTUNITY"] else "🤝 Existing Partner"
        ot = titles_by_org.get(org, {})
        pat, pub = ot.get("Patents", []), ot.get("Publications", [])
        tl = ""
        if pat:
            tl += f"  Sample patent titles: {'; '.join(pat[:3])}\n"
        if pub:
            tl += f"  Sample publication titles: {'; '.join(pub[:3])}\n"
        focus_pct = int(round(float(r["FOCUS"]) * 100))
        recent_pct = int(round(float(r["RECENT"]) * 100))
        # existing-partner queries lead with the literal collaboration count (the ranking basis);
        # recommendations lead with the composite match score.
        headline = (f"  Joint works with NUS in this field: {int(r['FCOLLAB'])}\n" if existing_only
                    else f"  Match score: {int(r['MATCH_SCORE'])}/100\n")
        rows.append(
            f"- {org} ({r['ORG_CATEGORY']}) [{tier}]\n"
            f"{headline}"
            f"  Field activity: {int(r['FIELD_CNT'])} patents/publications in this field\n"
            f"  Research focus: {focus_pct}% of their total output is in this field\n"
            f"  Recency: {recent_pct}% of their field work is from the last 3 years\n"
            f"  Academic track record: co-published / co-filed with {int(r['N_INST'])} distinct research institutes\n"
            f"{tl}"
        )
    data_str = "\n".join(rows)
    subject_context = f" in {subject}" if subject else ""
    mode_desc = ("existing industry partners actively collaborating with NUS, ranked by number of joint works in this field"
                 if existing_only else "recommended industry partners for NUS, ranked by match score")
    signals_extra = ("- Existing collaboration: [describe the current NUS relationship in this field]\n"
                     if existing_only else "")
    headline_metric = ("[N] joint works with NUS" if existing_only else "Match [score]/100")

    prompt = f"""You are a research collaboration advisor at the National University of Singapore (NUS).

Write structured recommendations for the top {len(recs_df)} {mode_desc}{subject_context}.

CRITICAL DEFINITIONS — read carefully before writing:
- "Match score" (0–100) is a composite ranking of collaboration potential for THIS field; higher = stronger fit, 100 = best-matched partner in this search. It is NOT a count of anything.
- "Field activity" = how many patents/publications the organisation has in this subject (their own output).
- "Research focus" = the share of their entire portfolio that is in this field (specialists score high).
- "Recency" = the share of their field work from the last 3 years (currently-active orgs score high).
- "Academic track record" = how many research institutes they have co-published or co-filed patents with — a proxy for willingness to work with academia.
- 🆕 New Opportunity = NO prior direct collaboration with NUS. Do NOT imply any existing joint work — frame overlap purely as unexplored potential.
- 🤝 Existing Partner = already collaborates with NUS; you may reference the depth of the relationship.

Use the tier labels exactly as shown. Be specific and data-driven; cite the sample titles as evidence of what they actually work on.
Write in a professional but accessible tone for senior stakeholders.
Do NOT add any title or heading before the recommendations. Start directly with the first --- divider.
Use only **bold** for emphasis — no # or ## headings anywhere.

Data:
{data_str}

Format each recommendation exactly as follows (markdown):

---
**[Rank]. [Organisation Name]** — [tier label] · {headline_metric}

**About:** 1-2 sentences on what the organisation does and their research focus, citing specific titles as evidence.

**Why they rank here:**
- Field activity: [N works in this field] — describe their footprint
- Focus & recency: [how concentrated in this field, and how current their work is]
- Academic track record: [their history of collaborating with research institutes]
{signals_extra.rstrip()}
**Why collaborate:** 1-2 sentences on the strategic rationale, framed for the tier (new = untapped potential; existing = deepen or expand).

**Strategic note:** One sentence on the specific next step this partner represents.
"""
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=min(8000, 500 + len(recs_df) * 400),
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


def _partner_records_cte(partner: str) -> str:
    """Shared CTE: normalized names + records (UIDs) for a parent brand."""
    p = partner.replace("'", "''")
    return f"""
WITH pnames AS (
    SELECT DISTINCT TRIM(UPPER(n.VALUE::STRING)) AS NNAME
    FROM INDUSTRY_AGG.PUBLIC.ENTITIES,
         LATERAL FLATTEN(INPUT => SPLIT(
            COALESCE(NORMALIZED_NAMES_PAT,'') || '|' || COALESCE(NORMALIZED_NAMES_PUB,''), '|')) n
    WHERE PARENT_BRAND = '{p}' AND TRIM(n.VALUE::STRING) <> ''
),
p_uids AS (
    SELECT DISTINCT p.UID
    FROM INDUSTRY_AGG.PUBLIC.PAT_PUB p,
         LATERAL FLATTEN(INPUT => SPLIT(p.NORMALIZED_NAMES_CONCAT, '|')) f
    JOIN pnames ON TRIM(UPPER(f.VALUE::STRING)) = NNAME
)"""


@st.cache_data(ttl=3600, show_spinner=False)
def flat_partner_subjects(partner: str) -> pd.DataFrame:
    """The partner's areas of focus (QS faculty areas), by record count."""
    sql = _partner_records_cte(partner) + """
SELECT TRIM(a.VALUE::STRING) AS AREA, COUNT(DISTINCT p.UID) AS CNT
FROM p_uids pu JOIN INDUSTRY_AGG.PUBLIC.PAT_PUB p ON p.UID = pu.UID,
     LATERAL FLATTEN(INPUT => SPLIT(p.QS_SUBJECT_AREA, '|')) a
WHERE TRIM(a.VALUE::STRING) NOT IN ('', '-')
GROUP BY 1 ORDER BY 2 DESC
"""
    return run_query(sql)


@st.cache_data(ttl=3600, show_spinner=False)
def flat_partner_nus_units(partner: str) -> pd.DataFrame:
    """NUS units the partner has co-published / co-filed with."""
    sql = _partner_records_cte(partner) + """
SELECT TRIM(u.VALUE::STRING) AS UNIT, COUNT(DISTINCT p.UID) AS CNT
FROM p_uids pu JOIN INDUSTRY_AGG.PUBLIC.PAT_PUB p ON p.UID = pu.UID,
     LATERAL FLATTEN(INPUT => SPLIT(p.UNITS, '|')) u
WHERE p.NUS_IP = TRUE AND p.UNITS IS NOT NULL AND TRIM(u.VALUE::STRING) <> ''
GROUP BY 1 ORDER BY 2 DESC
"""
    return run_query(sql)


@st.cache_data(ttl=3600, show_spinner=False)
def flat_partner_collaborators(partner: str, limit: int = 15) -> pd.DataFrame:
    """Other organisations (parent brands, excl. NUS) the partner co-appears with."""
    p = partner.replace("'", "''")
    sql = _partner_records_cte(partner) + f""",
comap AS (
    SELECT DISTINCT TRIM(UPPER(n.VALUE::STRING)) AS NNAME, PARENT_BRAND, CATEGORY, NUS_AFFILIATED
    FROM INDUSTRY_AGG.PUBLIC.ENTITIES,
         LATERAL FLATTEN(INPUT => SPLIT(
            COALESCE(NORMALIZED_NAMES_PAT,'') || '|' || COALESCE(NORMALIZED_NAMES_PUB,''), '|')) n
    WHERE TRIM(n.VALUE::STRING) <> ''
)
SELECT c.PARENT_BRAND AS ORG, ANY_VALUE(c.CATEGORY) AS CATEGORY, COUNT(DISTINCT p.UID) AS CNT
FROM p_uids pu JOIN INDUSTRY_AGG.PUBLIC.PAT_PUB p ON p.UID = pu.UID,
     LATERAL FLATTEN(INPUT => SPLIT(p.NORMALIZED_NAMES_CONCAT, '|')) f
JOIN comap c ON TRIM(UPPER(f.VALUE::STRING)) = c.NNAME
WHERE c.PARENT_BRAND <> '{p}' AND c.NUS_AFFILIATED = FALSE AND c.CATEGORY <> 'Individual'
GROUP BY 1 ORDER BY 3 DESC
LIMIT {int(limit)}
"""
    return run_query(sql)


# JS injected into PyVis HTML: click a node to focus it + dim the rest; click empty space to reset.
_DIM_JS = """
try {
  var _allNodes = nodes.get({returnType:"Object"});
  var _allEdges = edges.get({returnType:"Object"});
  var _orig = {}, _origFont = {}, _origEdge = {};
  for (var _id in _allNodes){
    _orig[_id] = _allNodes[_id].color;
    _origFont[_id] = (_allNodes[_id].font && _allNodes[_id].font.color) ? _allNodes[_id].font.color : '#ffffff';
  }
  for (var _eid in _allEdges){ _origEdge[_eid] = _allEdges[_eid].color; }
  var _active = false;
  function _pushUpdates(){
    var un=[]; for (var id in _allNodes){ un.push(_allNodes[id]); } nodes.update(un);
    var ue=[]; for (var eid in _allEdges){ ue.push(_allEdges[eid]); } edges.update(ue);
  }
  network.on("click", function(params){
    if (params.nodes.length > 0){
      _active = true;
      var sel = params.nodes[0];
      var keep = network.getConnectedNodes(sel); keep.push(sel);
      var keepE = network.getConnectedEdges(sel);
      for (var id in _allNodes){
        if (keep.indexOf(id) === -1){
          _allNodes[id].color = 'rgba(150,150,150,0.10)';
          _allNodes[id].font  = {color: 'rgba(210,210,210,0.15)'};
        } else {
          _allNodes[id].color = _orig[id];
          _allNodes[id].font  = {color: _origFont[id]};
        }
      }
      var keepEStr = keepE.map(String);
      for (var eid in _allEdges){
        _allEdges[eid].color = (keepEStr.indexOf(String(eid)) === -1) ? 'rgba(150,150,150,0.05)' : _origEdge[eid];
      }
    } else if (_active){
      _active = false;
      for (var id in _allNodes){
        _allNodes[id].color = _orig[id];
        _allNodes[id].font  = {color: _origFont[id]};
      }
      for (var eid in _allEdges){ _allEdges[eid].color = _origEdge[eid]; }
    }
    _pushUpdates();
  });
} catch(e) {}
// centre the graph in the viewport. Re-fit on the first few draws (catches iframe size-settling
// for the default/visible tab) and on resize, so every tab centres regardless of load order.
try {
  function _doFit(){ try { network.fit({animation:false}); } catch(e) {} }
  var _fitCount = 0;
  network.on("afterDrawing", function(){ if (_fitCount < 6){ _fitCount++; _doFit(); } });
  network.on("resize", _doFit);
  network.once("stabilizationIterationsDone", _doFit);
  network.once("stabilized", _doFit);
  setTimeout(_doFit, 400);
  setTimeout(_doFit, 1500);
} catch(e) {}
"""


@st.cache_data(ttl=3600, show_spinner=False)
def build_partner_flat_graph(partner: str, view: str = "collaborators") -> str:
    """PyVis graph centred on a partner. `view` ∈ {'subjects','units','collaborators'}."""
    net = Network(height="540px", width="100%", bgcolor="#1a1a1a",
                  font_color="#ffffff", directed=False, notebook=False, cdn_resources="in_line")
    net.force_atlas_2based(gravity=-45, central_gravity=0.012, spring_length=140,
                           spring_strength=0.08, damping=0.4, overlap=0)
    net.add_node(partner, label=partner, color="#ff6b6b", size=34,
                 title=f"{partner} (recommended partner)", shape="dot")

    if view == "subjects":
        for _, r in flat_partner_subjects(partner).head(8).iterrows():
            nid = f"subj::{r['AREA']}"
            net.add_node(nid, label=r["AREA"].title(), color="#F4B183", size=16,
                         title=f"Focus area · {int(r['CNT'])} records", shape="square")
            net.add_edge(partner, nid, color="#F4B183", width=1 + (r["CNT"] ** 0.3))

    elif view == "units":
        units = flat_partner_nus_units(partner)
        if not units.empty:
            net.add_node("NUS_HUB", label="NUS", color="#ff9933", size=24,
                         title="National University of Singapore", shape="dot")
            net.add_edge(partner, "NUS_HUB", color="#ff9933", width=3)
            for _, r in units.head(20).iterrows():
                nid = f"unit::{r['UNIT']}"
                net.add_node(nid, label=r["UNIT"], color="#ffd27f", size=13,
                             title=f"NUS unit · {int(r['CNT'])} joint works", shape="triangle")
                net.add_edge("NUS_HUB", nid, color="#ffd27f", width=1 + (r["CNT"] ** 0.4))

    else:  # collaborators
        _catcol = {"Corporation": "#ccccff", "Institute": "#33cccc",
                   "Hospital": "#9DC3E6", "Government / Non-profit": "#c9a0dc"}
        for _, r in flat_partner_collaborators(partner, limit=25).iterrows():
            nid = f"org::{r['ORG']}"
            net.add_node(nid, label=r["ORG"], color=_catcol.get(r["CATEGORY"], "#D9D9D9"),
                         size=13, title=f"{r['ORG']} ({r['CATEGORY']}) · {int(r['CNT'])} shared works")
            net.add_edge(partner, nid, color="#888888", width=1 + (r["CNT"] ** 0.3))

    html = net.generate_html()
    return html.replace("return network;", _DIM_JS + "\n        return network;")


def run_recommendation_subject_edges(
    institution: str,
    org_ids: list,
    subject_filter: str = None,
) -> pd.DataFrame:
    """Fetch shared subject edges between recommended orgs and the institution's subjects."""
    if not org_ids:
        return pd.DataFrame()
    safe_inst = sql_escape(institution)
    subject_clause = (
        f"AND (SOURCE_NAME ILIKE '%{sql_escape(subject_filter)}%' OR TARGET_NAME ILIKE '%{sql_escape(subject_filter)}%')"
        if subject_filter else ""
    )
    quoted = ", ".join(f"'{sql_escape(str(oid))}'" for oid in org_ids)
    sql = f"""
WITH inst_subjects AS (
    SELECT DISTINCT
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET ELSE SOURCE END AS SUBJECT_ID,
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET_NAME ELSE SOURCE_NAME END AS SUBJECT_NAME
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE (SOURCE_NAME ILIKE '%{safe_inst}%' OR TARGET_NAME ILIKE '%{safe_inst}%')
    AND EDGE_TYPE IN ('Applicant_Subject', 'Institute_Subject')
    {subject_clause}
)
SELECT
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN TARGET ELSE SOURCE END AS ORG_ID,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN TARGET_NAME ELSE SOURCE_NAME END AS ORG_NAME,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN SOURCE ELSE TARGET END AS SUBJECT_ID,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN SOURCE_NAME ELSE TARGET_NAME END AS SUBJECT_NAME,
    IP_TYPE,
    WEIGHT
FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
WHERE EDGE_TYPE IN ('Applicant_Subject', 'Institute_Subject')
AND (SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) OR TARGET IN (SELECT SUBJECT_ID FROM inst_subjects))
AND (SOURCE IN ({quoted}) OR TARGET IN ({quoted}))
"""
    return run_query(sql)


def run_similar_no_collab_subject_edges(
    institution: str,
    org_ids: list,
    ip_type: str = None,
    subject_filter: str = None,
) -> pd.DataFrame:
    """
    For a list of matched orgs, fetch their actual connections to shared subjects
    so we can draw org → subject edges in the graph.
    """
    if not org_ids:
        return pd.DataFrame()

    safe_inst = sql_escape(institution)
    ip_filter = f"AND IP_TYPE = '{sql_escape(ip_type)}'" if ip_type and ip_type != "All" else ""

    if ip_type == "Patents":
        subject_edge_types = "('Applicant_Subject')"
    elif ip_type == "Publications":
        subject_edge_types = "('Institute_Subject')"
    else:
        subject_edge_types = "('Applicant_Subject', 'Institute_Subject')"

    subject_clause = (
        f"AND (SOURCE_NAME ILIKE '%{sql_escape(subject_filter)}%' OR TARGET_NAME ILIKE '%{sql_escape(subject_filter)}%')"
        if subject_filter else ""
    )

    # Quote org IDs for SQL IN clause
    quoted_ids = ", ".join(f"'{sql_escape(str(oid))}'" for oid in org_ids)

    sql = f"""
WITH inst_subjects AS (
    SELECT DISTINCT
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET ELSE SOURCE END AS SUBJECT_ID,
        CASE WHEN SOURCE_NAME ILIKE '%{safe_inst}%' THEN TARGET_NAME ELSE SOURCE_NAME END AS SUBJECT_NAME
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE (SOURCE_NAME ILIKE '%{safe_inst}%' OR TARGET_NAME ILIKE '%{safe_inst}%')
    AND EDGE_TYPE IN {subject_edge_types}
    {ip_filter}
    {subject_clause}
)
SELECT
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN TARGET ELSE SOURCE END AS ORG_ID,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN TARGET_NAME ELSE SOURCE_NAME END AS ORG_NAME,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN SOURCE ELSE TARGET END AS SUBJECT_ID,
    CASE WHEN SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) THEN SOURCE_NAME ELSE TARGET_NAME END AS SUBJECT_NAME,
    WEIGHT
FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
WHERE EDGE_TYPE IN {subject_edge_types}
{ip_filter}
AND (SOURCE IN (SELECT SUBJECT_ID FROM inst_subjects) OR TARGET IN (SELECT SUBJECT_ID FROM inst_subjects))
AND (SOURCE IN ({quoted_ids}) OR TARGET IN ({quoted_ids}))
"""
    return run_query(sql)


@st.cache_data(ttl=3600, show_spinner=False)
def build_similar_no_collab_graph(results_df: pd.DataFrame, edges_df: pd.DataFrame) -> str:
    """
    Build a bipartite graph: organisations (blue) on the left, shared subject areas (orange) on the right.
    Edges connect orgs to their matching subjects. Node size = number of connections.
    """
    net = Network(
        height="750px",
        width="100%",
        bgcolor="#1a1a1a",
        font_color="#ffffff",
        directed=False,
        notebook=False,
        cdn_resources="in_line",
    )
    net.force_atlas_2based(
        gravity=-50,
        central_gravity=0.01,
        spring_length=100,
        spring_strength=0.08,
        damping=0.4,
        overlap=0,
    )

    added_orgs = set()
    added_subjects = set()

    # Build lookup for org metadata from results_df
    org_meta = {
        str(row["ORG_ID"]): {
            "name": str(row["ORG_NAME"]),
            "category": str(row.get("ORG_CATEGORY", "")),
            "shared": int(row["SHARED_SUBJECTS"]),
            "weight": float(row["TOTAL_WEIGHT"]),
        }
        for _, row in results_df.iterrows()
    }

    for _, row in edges_df.iterrows():
        org_id = str(row["ORG_ID"])
        org_name = str(row["ORG_NAME"])
        subj_id = str(row["SUBJECT_ID"])
        subj_name = str(row["SUBJECT_NAME"])
        weight = float(row["WEIGHT"])

        meta = org_meta.get(org_id, {})

        if org_id not in added_orgs:
            shared = meta.get("shared", 1)
            net.add_node(
                org_id,
                label=org_name,
                title=f"{org_name}\nCategory: {meta.get('category', '—')}\nShared subjects: {shared}\nTotal strength: {meta.get('weight', 0):.0f}",
                color="#9DC3E6",
                value=shared,
            )
            added_orgs.add(org_id)

        if subj_id not in added_subjects:
            net.add_node(
                subj_id,
                label=subj_name,
                title=f"{subj_name}\nResearch subject area",
                color="#F4B183",
                value=3,
            )
            added_subjects.add(subj_id)

        net.add_edge(
            org_id,
            subj_id,
            value=weight,
            title=f"Strength: {weight:.0f}",
        )

    return net.generate_html(notebook=False)


def inject_layout_controls(html: str) -> str:
    """Replace physics panel with clean preset buttons that don't interfere with scrolling."""
    controls = """
    <div style="padding:8px 12px; display:flex; gap:8px; flex-wrap:wrap; background:#1a1a1a; border-bottom:1px solid #333;">
        <button onclick="setForceAtlas()" style="background:#333;color:#fff;border:1px solid #555;padding:5px 14px;border-radius:4px;cursor:pointer;font-size:12px;">🔄 ForceAtlas2</button>
        <button onclick="setBarnesHut()" style="background:#333;color:#fff;border:1px solid #555;padding:5px 14px;border-radius:4px;cursor:pointer;font-size:12px;">🌐 Barnes-Hut</button>
        <button onclick="freezeGraph()" style="background:#333;color:#fff;border:1px solid #555;padding:5px 14px;border-radius:4px;cursor:pointer;font-size:12px;">❄️ Freeze</button>
        <button onclick="network.fit()" style="background:#333;color:#fff;border:1px solid #555;padding:5px 14px;border-radius:4px;cursor:pointer;font-size:12px;">⊡ Fit to Screen</button>
    </div>
    <script>
    function setForceAtlas() {
        network.setOptions({physics:{enabled:true,solver:'forceAtlas2Based',forceAtlas2Based:{gravity:-50,centralGravity:0.01,springLength:100,springConstant:0.08,damping:0.4,overlap:0}}});
    }
    function setBarnesHut() {
        network.setOptions({physics:{enabled:true,solver:'barnesHut',barnesHut:{gravitationalConstant:-30000,centralGravity:0.3,springLength:180,springConstant:0.02,damping:0.8,avoidOverlap:0.5}}});
    }
    function freezeGraph() {
        network.setOptions({physics:{enabled:false}});
    }
    </script>
    """
    return html.replace("<div id=\"mynetwork\"", controls + "<div id=\"mynetwork\"")


def inject_png_download(html: str, filename: str = "graph.png") -> str:
    """Inject a PNG download button into PyVis HTML using canvas capture."""
    button_js = f"""
    <div style="text-align:right; padding: 6px 12px;">
      <button onclick="downloadPNG()" style="
        background:#0068C9; color:white; border:none;
        padding:7px 16px; border-radius:6px; cursor:pointer; font-size:13px;">
        ⬇️ Download PNG
      </button>
    </div>
    <script>
    function downloadPNG() {{
      var canvas = document.querySelector('canvas');
      if (!canvas) {{ alert('Graph not ready yet — please wait a moment and try again.'); return; }}
      var link = document.createElement('a');
      link.download = '{filename}';
      link.href = canvas.toDataURL('image/png');
      link.click();
    }}
    </script>
    """
    # Inject just before closing </body>
    return html.replace("</body>", button_js + "</body>")


def _add_mixed_run(paragraph, text: str):
    """Add text to a paragraph, rendering **bold** markdown segments as bold runs."""
    for i, part in enumerate(re.split(r'(\*\*[^*]+\*\*)', text)):
        run = paragraph.add_run(re.sub(r'\*\*', '', part))
        if part.startswith("**") and part.endswith("**"):
            run.bold = True


def build_recommendation_docx(rec_text: str, institution: str, subject_filter: str = None) -> bytes:
    """Convert the AI recommendation markdown text into a formatted .docx file."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    subject_context = f" — {subject_filter.title()}" if subject_filter else ""

    title_para = doc.add_paragraph()
    tr = title_para.add_run("Research Collaboration Recommendations")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_para = doc.add_paragraph()
    sr = sub_para.add_run(f"{institution.title()}{subject_context}")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    blocks = [b.strip() for b in rec_text.split("---") if b.strip()]
    for idx, block in enumerate(blocks):
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Org heading: **N. Name** — 🆕/🤝 tier
            if re.match(r'^\*\*.+\*\*\s*[—–]', line) and ("🆕" in line or "🤝" in line):
                p = doc.add_paragraph()
                run = p.add_run(re.sub(r'\*\*', '', line))
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            # Bullet point
            elif line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                _add_mixed_run(p, line[2:])

            # Section label: **Label:** optional inline content
            elif re.match(r'^\*\*[^*]+:\*\*', line):
                m = re.match(r'^\*\*([^*]+):\*\*\s*(.*)', line)
                if m:
                    p = doc.add_paragraph()
                    lr = p.add_run(f"{m.group(1)}: ")
                    lr.bold = True
                    lr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                    if m.group(2):
                        p.add_run(m.group(2))

            else:
                p = doc.add_paragraph()
                _add_mixed_run(p, line)

        if idx < len(blocks) - 1:
            doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@st.cache_data(ttl=3600, show_spinner=False)
def build_pyvis_graph(df: pd.DataFrame, highlight_term: str = None) -> str:
    net = Network(
        height="750px",
        width="100%",
        bgcolor="#1a1a1a",
        font_color="#ffffff",
        directed=False,
        notebook=False,
        cdn_resources="in_line",
    )

    net.force_atlas_2based(
        gravity=-50,
        central_gravity=0.01,
        spring_length=100,
        spring_strength=0.08,
        damping=0.4,
        overlap=0,
    )

    added_nodes = set()
    term = highlight_term.upper() if highlight_term else None

    for _, row in df.iterrows():
        source_id = row["SOURCE"]
        target_id = row["TARGET"]
        source_name = row["SOURCE_NAME"]
        target_name = row["TARGET_NAME"]
        source_type = row["SOURCE_TYPE"]
        target_type = row["TARGET_TYPE"]
        weight = row["WEIGHT"]
        edge_type = row["EDGE_TYPE"]

        # Determine highlight state
        source_match = term and term in str(source_name).upper()
        target_match = term and term in str(target_name).upper()
        edge_highlighted = source_match or target_match

        if source_id not in added_nodes:
            base_color = get_node_color(source_type, source_name, row["SOURCE_NUS_AFFILIATED"])
            if term:
                color = base_color if source_match else "#E8E8E8"
                border = "#FF6600" if source_match else "#E8E8E8"
                node_size = 20 if source_match else 5
            else:
                color = base_color
                border = base_color
                node_size = 10

            net.add_node(
                source_id,
                label=source_name if (not term or source_match) else "",
                title=f"{source_name}\nType: {source_type}\nCategory: {row['SOURCE_CATEGORY']}",
                color={"background": color, "border": border},
                value=node_size,
            )
            added_nodes.add(source_id)

        if target_id not in added_nodes:
            base_color = get_node_color(target_type, target_name, row["TARGET_NUS_AFFILIATED"])
            if term:
                color = base_color if target_match else "#E8E8E8"
                border = "#FF6600" if target_match else "#E8E8E8"
                node_size = 20 if target_match else 5
            else:
                color = base_color
                border = base_color
                node_size = 10

            net.add_node(
                target_id,
                label=target_name if (not term or target_match) else "",
                title=f"{target_name}\nType: {target_type}\nCategory: {row['TARGET_CATEGORY']}",
                color={"background": color, "border": border},
                value=node_size,
            )
            added_nodes.add(target_id)

        net.add_edge(
            source_id,
            target_id,
            value=float(weight) if edge_highlighted or not term else 0.1,
            color="#FF6600" if edge_highlighted else "#DDDDDD",
            title=f"Connection type: {edge_type}\nStrength: {weight}",
        )

    html = net.generate_html(notebook=False)
    return html


# -----------------------------
# LLM chat helper
# -----------------------------
SYSTEM_PROMPT = """You are a research collaboration discovery assistant helping users explore a graph network of patents and academic publications.

Based on the user's natural language request, first decide if they are:
A) Asking to filter/explore the graph network
B) Asking a general question (about an industry partner, institution, research topic, or concept)

Return a JSON object with the following fields:

{
  "response_type": "<string>",          // REQUIRED. One of:
                                        // "graph_query"    — user wants to filter or explore the graph network visually
                                        //   (e.g. "show the network", "show connections", "show the graph", "visualise").
                                        // "general_answer" — user wants information about an industry partner, institution, topic, or concept
                                        // "recommendation" — user wants a ranked list of partners for an institution.
                                        //   Use this for BOTH existing collaborators AND potential new ones.
                                        //   Set existing_only=true if user asks who already collaborates or collaborates most.
                                        //   Set existing_only=false if user asks for recommendations, suggestions, or potential partners.
  "answer": "<string or null>",         // ONLY for general_answer: a helpful, concise answer (2-4 paragraphs).
                                        // Include: what the org does, their main research/business areas,
                                        // why they might be a good collaboration partner, and any notable facts.
                                        // For graph_query, set this to null.
  "query_mode": "<string>",             // for graph_query only. One of:
                                        // "standard"          — normal graph filter mode (default)
                                        // "similar_no_collab" — find orgs with similar research interests that have NOT yet collaborated with the searched institution.
  "ip_type": "<string or null>",        // "Patents" or "Publications", or null for both. Available: [AVAILABLE_IP_TYPES]
  "edge_type": "<string or null>",      // type of connection, or null for all. Available: [AVAILABLE_EDGE_TYPES]
                                        // Applicant_Applicant = patent co-applicants
                                        // Applicant_Subject = patent-to-subject links
                                        // Institute_Subject = publication-to-subject links
                                        // Institution_Institution = publication co-authors
  "search_term": "<string or null>",    // institution name to focus on, or null
  "category": "<string or null>",       // organisation category, or null for all. Available: [AVAILABLE_CATEGORIES]
                                        // Leave null unless the user explicitly names a specific org type.
                                        // Do NOT set "Corporation" just because the user says "industry partners" — leave null.
  "min_weight": <integer or null>,      // minimum collaboration strength, or null to keep current
  "max_edges": <integer or null>,       // max edges to load (20–1000), or null to keep current.
                                        // For "top N" requests: set max_edges=300 and top_n_nodes=N instead.
  "top_n_nodes": <integer or null>,     // for "top N partners" in standard mode, set to N. Leave null otherwise.
  "top_n_results": <integer or null>,   // for "similar_no_collab" and "recommendation" mode: how many results to return.
                                        // Leave null unless the user asks for a specific number (e.g. "top 5").
  "subject_filter": "<string or null>", // scope the search to QS subject area(s). Each subject MUST match one of: [AVAILABLE_SUBJECTS]
                                        // Map natural language to the QS subject name(s), e.g.:
                                        //   "AI", "artificial intelligence", "machine learning" → "COMPUTER SCIENCE & INFORMATION SYSTEMS"
                                        //   "biomedical", "life sciences" → "BIOLOGICAL SCIENCES"
                                        //   "engineering" → the most specific match e.g. "ENGINEERING - ELECTRICAL & ELECTRONIC"
                                        // For "recommendation" queries ONLY, a cross-disciplinary topic may map to
                                        // MULTIPLE subjects joined by " | " (max 3), e.g.:
                                        //   "fintech" → "COMPUTER SCIENCE & INFORMATION SYSTEMS | ACCOUNTING & FINANCE"
                                        //   "biotech" → "BIOLOGICAL SCIENCES | MEDICINE | CHEMISTRY"
                                        // For graph_query, always use a SINGLE subject.
                                        // If ambiguous, pick the closest match or leave null for all subjects.
  "existing_only": <boolean>,            // for "recommendation" mode only.
                                        // true  = user wants EXISTING collaborators only ("collaborate most", "top collaborators", "who works with NUS")
                                        // false = user wants potential/suggested partners ("recommend", "suggest", "potential partners")
  "explanation": "<friendly 1-2 sentence explanation of what the results will show, or null for general_answer. Always use 'industry partners' — never 'corporations', 'companies', or 'firms'.>"
}

Rules:
- response_type is ALWAYS required.
- For "general_answer": fill "answer" with a helpful response, set all filter fields to null.
- For "graph_query": fill filter fields as needed, set "answer" to null.
- query_mode defaults to "standard" for graph_query unless user clearly wants "similar_no_collab".
- ip_type must exactly match one of: [AVAILABLE_IP_TYPES]
- edge_type must exactly match one of: [AVAILABLE_EDGE_TYPES]
- category must exactly match one of: [AVAILABLE_CATEGORIES]
- subject_filter must exactly match one of: [AVAILABLE_SUBJECTS] — never invent a subject name
- If the user mentions "NUS" or "National University of Singapore", set search_term to "NATIONAL UNIVERSITY OF SINGAPORE".
- If the user says "reset", "clear", "start over", or "show everything", set response_type="graph_query", query_mode="standard", ip_type=null, edge_type=null, search_term=null, min_weight=1, max_edges=200, category=null, top_n_nodes=null, top_n_results=null.
- For "similar_no_collab" mode, search_term is required.
- Always return ONLY valid JSON. No markdown fences, no extra text."""


def extract_filters_from_llm(
    user_message: str,
    chat_history: list,
    available_ip_types: list,
    available_edge_types: list,
    available_categories: list,
    available_subjects: list,
) -> dict: 

    system = SYSTEM_PROMPT.replace(
        "[AVAILABLE_IP_TYPES]", ", ".join(available_ip_types)
    ).replace(
        "[AVAILABLE_EDGE_TYPES]", ", ".join(available_edge_types)
    ).replace(
        "[AVAILABLE_CATEGORIES]", ", ".join(available_categories)
    ).replace(
        "[AVAILABLE_SUBJECTS]", ", ".join(available_subjects)
    )

    messages = []
    for turn in chat_history:
        messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": user_message})

    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=1024,
        system=system,
        messages=messages,
    )

    raw = response.content[0].text.strip()

    # Strip markdown fences if model adds them anyway
    raw = re.sub(r"^```(?:json)?", "", raw).strip()
    raw = re.sub(r"```$", "", raw).strip()

    return json.loads(raw)


def apply_llm_filters(parsed: dict, current_state: dict, available_ip_types: list, available_edge_types: list, available_categories: list) -> dict:
    """Merge LLM-extracted filters onto the current filter state."""
    new_state = current_state.copy()

    if parsed.get("ip_type") is not None:
        val = parsed["ip_type"].upper().rstrip("S")  # "PATENTS"→"PATENT", "PUBLICATIONS"→"PUBLICATION"
        match = next((t for t in available_ip_types if t.upper().rstrip("S") == val), None)
        new_state["ip_type"] = match if match else "All"

    if parsed.get("edge_type") is not None:
        val = parsed["edge_type"].upper().replace("-", "_").replace(" ", "_")
        match = next((t for t in available_edge_types if t.upper().replace("-", "_") == val), None)
        new_state["edge_type"] = match if match else "All"

    if parsed.get("category") is not None:
        val = parsed["category"].upper()
        match = next((c for c in available_categories if c.upper() == val), None)
        new_state["category"] = match if match else "All"

    if parsed.get("search_term") is not None:
        new_state["search_term"] = parsed["search_term"]

    if parsed.get("min_weight") is not None:
        new_state["min_weight"] = max(1, int(parsed["min_weight"]))

    if parsed.get("max_edges") is not None:
        new_state["max_edges"] = max(20, min(1000, int(parsed["max_edges"])))

    new_state["top_n_nodes"] = int(parsed["top_n_nodes"]) if parsed.get("top_n_nodes") is not None else None
    new_state["query_mode"] = parsed.get("query_mode", "standard") or "standard"
    new_state["top_n_results"] = int(parsed["top_n_results"]) if parsed.get("top_n_results") is not None else None
    new_state["subject_filter"] = parsed.get("subject_filter") or None

    return new_state


# -----------------------------
# Load filter options
# -----------------------------
edge_types_df = run_query("""
    SELECT DISTINCT EDGE_TYPE
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE EDGE_TYPE IS NOT NULL
    ORDER BY EDGE_TYPE
""")

ip_types_df = run_query("""
    SELECT DISTINCT IP_TYPE
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE IP_TYPE IS NOT NULL
    ORDER BY IP_TYPE
""")

categories_df = run_query("""
    SELECT DISTINCT SOURCE_CATEGORY AS CATEGORY
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE SOURCE_CATEGORY IS NOT NULL
    UNION
    SELECT DISTINCT TARGET_CATEGORY AS CATEGORY
    FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
    WHERE TARGET_CATEGORY IS NOT NULL
    ORDER BY CATEGORY
""")

edge_types_raw = edge_types_df["EDGE_TYPE"].tolist()
ip_types_raw = ip_types_df["IP_TYPE"].tolist()
categories_raw = categories_df["CATEGORY"].tolist()

subjects_df = run_query("""
    SELECT DISTINCT NODE_NAME AS SUBJECT_NAME
    FROM GRAPH_NETWORK.GRAPH.ALL_NODES
    WHERE NODE_TYPE = 'Subject'
    AND NODE_NAME != 'NO SUBJECT DETECTED'
    ORDER BY SUBJECT_NAME
""")
subjects_raw = subjects_df["SUBJECT_NAME"].tolist()

edge_types = ["All"] + edge_types_raw
ip_types = ["All"] + ip_types_raw
categories = ["All"] + categories_raw


# -----------------------------
# Session state initialisation
# -----------------------------
if "filter_state" not in st.session_state:
    st.session_state.filter_state = {
        "ip_type": "All",
        "edge_type": "All",
        "search_term": "",
        "min_weight": 1,
        "max_edges": 200,
        "category": "All",
        "top_n_nodes": None,
        "query_mode": "standard",
        "top_n_results": None,
        "subject_filter": None,
        "has_queried": False,
        "recommendation_data": None,
    }

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []  # list of {"role": ..., "content": ...}

if "chat_display" not in st.session_state:
    st.session_state.chat_display = []  # list of {"role": ..., "content": ..., "filters": ...}


# -----------------------------
# Sidebar: manual filters + chat
# -----------------------------
with st.sidebar:
    st.header("⚙️ Manual Filters")

    st.selectbox(
        "Research output type",
        ip_types,
        index=ip_types.index(st.session_state.filter_state["ip_type"]) if st.session_state.filter_state["ip_type"] in ip_types else 0,
        key="sb_ip_type",
        on_change=lambda: (st.session_state.filter_state.update({"ip_type": st.session_state.sb_ip_type, "top_n_nodes": None, "query_mode": "standard", "has_queried": True, "recommendation_data": None}), run_query.clear()),
    )

    subject_options = ["All"] + subjects_raw
    current_subject = st.session_state.filter_state.get("subject_filter") or "All"
    st.selectbox(
        "Research subject",
        subject_options,
        index=subject_options.index(current_subject) if current_subject in subject_options else 0,
        key="sb_subject_filter",
        on_change=lambda: (st.session_state.filter_state.update({"subject_filter": st.session_state.sb_subject_filter if st.session_state.sb_subject_filter != "All" else None, "top_n_nodes": None, "query_mode": "standard", "has_queried": True}), run_query.clear()),
    )

    st.selectbox(
        "Organisation category",
        categories,
        index=categories.index(st.session_state.filter_state["category"]) if st.session_state.filter_state["category"] in categories else 0,
        key="sb_category",
        on_change=lambda: (st.session_state.filter_state.update({"category": st.session_state.sb_category, "top_n_nodes": None, "query_mode": "standard", "has_queried": True}), run_query.clear()),
    )

    st.text_input(
        "Search for an institution or organisation",
        value=st.session_state.filter_state["search_term"],
        placeholder="e.g. NATIONAL UNIVERSITY OF SINGAPORE",
        key="sb_search",
        on_change=lambda: (st.session_state.filter_state.update({"search_term": st.session_state.sb_search, "top_n_nodes": None, "query_mode": "standard", "has_queried": True}), run_query.clear()),
    )

    st.number_input(
        "Minimum collaboration strength",
        min_value=1,
        value=st.session_state.filter_state["min_weight"],
        key="sb_min_weight",
        on_change=lambda: (st.session_state.filter_state.update({"min_weight": st.session_state.sb_min_weight, "has_queried": True}), run_query.clear()),
    )

    st.slider(
        "Maximum connections to load",
        min_value=20,
        max_value=1000,
        value=st.session_state.filter_state["max_edges"],
        step=20,
        key="sb_max_edges",
        on_change=lambda: (st.session_state.filter_state.update({"max_edges": st.session_state.sb_max_edges, "top_n_nodes": None, "has_queried": True}), run_query.clear()),
    )

    st.divider()

# -----------------------------
# Read effective filters (may have been updated by LLM)
# -----------------------------
fs = st.session_state.filter_state
selected_ip_type = fs["ip_type"]
selected_edge_type = fs["edge_type"]
selected_category = fs["category"]
search_term = fs["search_term"]
min_weight = fs["min_weight"]
max_edges = fs["max_edges"]
top_n_nodes = fs.get("top_n_nodes")
query_mode = fs.get("query_mode", "standard")
top_n_results = fs.get("top_n_results") or 20
subject_filter = fs.get("subject_filter")
has_queried = fs.get("has_queried", False)


# -----------------------------
# Two column layout
# -----------------------------
chat_col, graph_col = st.columns([2, 3], gap="large")

with graph_col:
    st.subheader("🗺️ Collaboration Network")
    st.markdown(
        "<span style='font-size:15px'>"
        "<span style='color:#ff9933'>■</span> NUS-affiliated &nbsp;|&nbsp; "
        "<span style='color:#ff6b6b'>■</span> New opportunity &nbsp;|&nbsp; "
        "<span style='color:#ccccff'>■</span> Patent applicant &nbsp;|&nbsp; "
        "<span style='color:#9DC3E6'>■</span> Existing partner &nbsp;|&nbsp; "
        "<span style='color:#33cccc'>■</span> Publication institute &nbsp;|&nbsp; "
        "<span style='color:#F4B183'>■</span> Shared Subject Area &nbsp;|&nbsp; "
        "<span style='color:#D9D9D9'>■</span> Other"
        "</span>"
        "<br><span style='font-size:13px; color:#aaaaaa'>Node size = number of shared subjects &nbsp;|&nbsp; Line thickness = collaboration strength</span>",
        unsafe_allow_html=True,
    )

    if not has_queried:
        st.info(
            "👋 **Welcome to the Potential Collaborator Finder!**\n\n"
            "Use the AI assistant on the left to get started. Here are some things you can ask:\n\n"
            "- _'Recommend industry partners for NUS in computer science'_\n"
            "- _'Which industry partners collaborate most with NUS in semiconductors?'_\n"
            "- _'Show industry partners with similar research interests to NUS that haven\\'t collaborated before'_\n\n"
            "Or use the manual filters in the sidebar."
        )

    elif query_mode == "recommendation":
        # --- Recommendation mode ---
        rec_data = fs.get("recommendation_data")
        if not rec_data:
            st.info("Ask the AI assistant for recommendations — try: _'Recommend industry partners for NUS in AI'_")
        elif rec_data.get("flat"):
            # ── Flat-model recommendation view (new scoring engine) ──
            recs_df = pd.DataFrame(rec_data["recs_df"])
            institution = rec_data["institution"]
            subject_filter = rec_data.get("subject_filter")
            existing_only = rec_data.get("existing_only", False)
            subject_context = f" in {subject_filter}" if subject_filter else ""
            if existing_only:
                st.markdown(f"Top industry partners **actively collaborating** with **NUS**{subject_context}, ranked by number of joint works in this field.")
            else:
                st.markdown(f"Recommended industry partners for **NUS**{subject_context}, ranked by match score.")

            st.subheader("📋 Summary")
            _focus = (recs_df["FOCUS"] * 100).round().astype(int).astype(str) + "%"
            _recent = (recs_df["RECENT"] * 100).round().astype(int).astype(str) + "%"
            if existing_only:
                # "collaborate most" — lead with the literal joint-works count (the ranking basis)
                _sdf = pd.DataFrame({
                    "Organisation": recs_df["ORG_NAME"].values,
                    "Category": recs_df["ORG_CATEGORY"].values,
                    "Joint works with NUS": recs_df["FCOLLAB"].values,
                    "Patents & publications in this field": recs_df["FIELD_CNT"].values,
                    "Recently active (Past 3 Years)": _recent.values,
                    "Works with universities": recs_df["N_INST"].values,
                })
            else:
                _status = recs_df["IS_NEW_OPPORTUNITY"].map(lambda x: "🆕 New" if x else "🤝 Existing")
                _sdf = pd.DataFrame({
                    "Organisation": recs_df["ORG_NAME"].values,
                    "Category": recs_df["ORG_CATEGORY"].values,
                    "Opportunity": _status.values,
                    "Match score": recs_df["MATCH_SCORE"].values,
                    "Patents & publications in this field": recs_df["FIELD_CNT"].values,
                    "Focused on this field": _focus.values,
                    "Recently active (Past 3 Years)": _recent.values,
                    "Works with universities": recs_df["N_INST"].values,
                })
            _sdf.index = range(1, len(_sdf) + 1)
            st.caption("Click a row to explore that partner's focus areas, NUS-unit collaborations, and other partners.")
            _colcfg = {
                "Match score": st.column_config.NumberColumn(
                    "Match score", help="Overall fit for this field, 0–100 (100 = best match in this search)."),
                "Opportunity": st.column_config.TextColumn(
                    "Opportunity", help="🆕 New = no prior NUS collaboration · 🤝 Existing = already works with NUS."),
                "Joint works with NUS": st.column_config.NumberColumn(
                    "Joint works with NUS", help="Patents/publications they've co-produced with NUS in this field."),
                "Patents & publications in this field": st.column_config.NumberColumn(
                    "Patents & publications in this field", help="How many of their patents/publications are in this research area."),
                "Focused on this field": st.column_config.TextColumn(
                    "Focused on this field", help="Share of their entire portfolio that is in this field — higher = more of a specialist."),
                "Recently active (Past 3 Years)": st.column_config.TextColumn(
                    "Recently active (Past 3 Years)", help="How recent their work in this field is — higher = more currently active."),
                "Works with universities": st.column_config.NumberColumn(
                    "Works with universities", help="Number of research institutes they've partnered with — a sign they're open to academic collaboration."),
            }
            _sel = st.dataframe(
                _sdf, on_select="rerun", selection_mode="single-row",
                use_container_width=True, column_config=_colcfg,
            )

            _rcol1, _rcol2, _rcol3 = st.columns([1, 1, 4])
            with _rcol1:
                st.download_button(
                    "⬇️ Download CSV", data=_sdf.to_csv().encode("utf-8"),
                    file_name="recommendations.csv", mime="text/csv", key="rec_flat_csv",
                )
            with _rcol2:
                _rt = rec_data.get("rec_text", "")
                if _rt:
                    _dx = build_recommendation_docx(_rt, institution, subject_filter)
                    st.download_button(
                        "⬇️ Download Report", data=_dx,
                        file_name="recommendations_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="rec_flat_docx",
                    )
            st.caption("Full write-up is shown in the chat panel on the left.")

            # ── Partner drill-down (click a row above) ──
            if _sel.selection.rows:
                _partner = str(recs_df.iloc[_sel.selection.rows[0]]["ORG_NAME"])
                st.markdown("---")
                st.subheader(f"🔎 {_partner.title()}")
                with st.spinner("Loading partner network…"):
                    _subj = flat_partner_subjects(_partner)
                    _units = flat_partner_nus_units(_partner)
                    _collab = flat_partner_collaborators(_partner, limit=15)

                with st.expander("🌐 Partner network — click a node to focus, click empty space to reset", expanded=True):
                    _gt1, _gt2, _gt3 = st.tabs(["🎯 Focus areas", "🏛 NUS units", "🤝 Other collaborators"])
                    with _gt1:
                        components.html(build_partner_flat_graph(_partner, "subjects"), height=560, scrolling=True)
                        st.markdown(
                            "<span style='font-size:13px'><span style='color:#ff6b6b'>■</span> Partner &nbsp;|&nbsp; "
                            "<span style='color:#F4B183'>■</span> Focus area</span>", unsafe_allow_html=True)
                    with _gt2:
                        if _units.empty:
                            st.info("No NUS-unit collaborations found — this partner hasn't co-published or co-filed with a NUS unit.")
                        else:
                            components.html(build_partner_flat_graph(_partner, "units"), height=560, scrolling=True)
                            st.markdown(
                                "<span style='font-size:13px'><span style='color:#ff6b6b'>■</span> Partner &nbsp;|&nbsp; "
                                "<span style='color:#ff9933'>■</span> NUS &nbsp;|&nbsp; "
                                "<span style='color:#ffd27f'>■</span> NUS unit</span>", unsafe_allow_html=True)
                    with _gt3:
                        components.html(build_partner_flat_graph(_partner, "collaborators"), height=560, scrolling=True)
                        st.markdown(
                            "<span style='font-size:13px'><span style='color:#ff6b6b'>■</span> Partner &nbsp;|&nbsp; "
                            "<span style='color:#33cccc'>■</span> Institute &nbsp;|&nbsp; "
                            "<span style='color:#ccccff'>■</span> Corporation &nbsp;|&nbsp; "
                            "<span style='color:#9DC3E6'>■</span> Hospital &nbsp;|&nbsp; "
                            "<span style='color:#c9a0dc'>■</span> Gov / Non-profit</span>", unsafe_allow_html=True)

                _da, _db = st.columns(2)
                with _da:
                    st.markdown("**Areas of focus**")
                    if not _subj.empty:
                        _sd = _subj.rename(columns={"AREA": "Area", "CNT": "Records"})
                        st.dataframe(_sd, use_container_width=True, hide_index=True)
                    else:
                        st.caption("—")
                with _db:
                    st.markdown("**NUS units collaborated with**")
                    if not _units.empty:
                        _ud = _units.rename(columns={"UNIT": "NUS Unit", "CNT": "Joint works"})
                        st.dataframe(_ud, use_container_width=True, hide_index=True)
                    else:
                        st.caption("No NUS-unit collaborations found.")

                st.markdown("**Other collaboration partners**")
                if not _collab.empty:
                    _cd = _collab.rename(columns={"ORG": "Organisation", "CATEGORY": "Category", "CNT": "Shared works"})
                    st.dataframe(_cd, use_container_width=True, hide_index=True)
                    st.download_button(
                        "⬇️ Download collaborators CSV", data=_cd.to_csv(index=False).encode("utf-8"),
                        file_name=f"{_partner[:30]}_collaborators.csv", mime="text/csv", key="partner_collab_csv",
                    )
                else:
                    st.caption("—")
    elif query_mode == "similar_no_collab":
        # --- Similar interests, no prior collaboration mode ---
        if not search_term.strip():
            st.warning("Please specify an institution to search from. Try asking: _'Show corporations with similar interests to NUS in publications that haven't collaborated with NUS'_")
        else:
            st.markdown(
                f"Showing organisations with **similar research interests** to **{search_term.title()}** "
                f"that have **not yet directly collaborated** with it. "
                "Ranked by number of shared subjects. Hover over nodes for details."
            )
            with st.spinner("Running analysis…"):
                similar_df = run_similar_no_collab_query(
                    institution=search_term.strip(),
                    ip_type=selected_ip_type if selected_ip_type != "All" else None,
                    category=selected_category if selected_category != "All" else None,
                    top_n=top_n_results,
                    subject_filter=subject_filter,
                )

            if similar_df.empty:
                st.info("No matches found. Try broadening the filters — e.g. remove the category filter or change the output type.")
            else:
                org_ids = similar_df["ORG_ID"].tolist()
                with st.spinner("Loading subject connections…"):
                    edges_df = run_similar_no_collab_subject_edges(
                        institution=search_term.strip(),
                        org_ids=org_ids,
                        ip_type=selected_ip_type if selected_ip_type != "All" else None,
                        subject_filter=subject_filter,
                    )

                # Derive shared subject names from edges and merge into similar_df
                if not edges_df.empty:
                    subj_names = (
                        edges_df.groupby("ORG_ID")["SUBJECT_NAME"]
                        .apply(lambda x: ", ".join(sorted(x.dropna().unique())))
                        .reset_index()
                        .rename(columns={"SUBJECT_NAME": "SHARED_SUBJECT_NAMES"})
                    )
                    similar_df = similar_df.merge(subj_names, on="ORG_ID", how="left")
                else:
                    similar_df["SHARED_SUBJECT_NAMES"] = ""

                # ── SUMMARY TABLE ─────────────────────────────────────────────
                st.subheader("📋 Ranked results")
                display_df = similar_df[[
                    "ORG_NAME", "ORG_CATEGORY",
                    "SHARED_SUBJECTS", "TOTAL_WEIGHT",
                    "SHARED_SUBJECT_NAMES",
                ]].copy()
                display_df.columns = [
                    "Organisation", "Category",
                    "Shared Subjects", "Patents & Publications",
                    "Shared Subject Areas",
                ]
                display_df.index = range(1, len(display_df) + 1)

                st.caption("Click a row to view that organisation's subject network.")
                sel = st.dataframe(
                    display_df,
                    on_select="rerun",
                    selection_mode="single-row",
                    use_container_width=True,
                )

                st.download_button(
                    label="⬇️ Download CSV",
                    data=display_df.to_csv().encode("utf-8"),
                    file_name="potential_partners.csv",
                    mime="text/csv",
                )

                # ── SELECTED ROW → FILTER GRAPH ───────────────────────────────
                selected_snc_name = None
                if sel.selection.rows:
                    idx = sel.selection.rows[0]
                    selected_row = similar_df.iloc[idx]
                    selected_snc_id = str(selected_row["ORG_ID"])
                    selected_snc_name = str(selected_row["ORG_NAME"])
                    filter_similar_df = similar_df[similar_df["ORG_ID"].astype(str) == selected_snc_id]
                    filter_edges_df = edges_df[edges_df["ORG_ID"].astype(str) == selected_snc_id]
                else:
                    filter_similar_df = similar_df
                    filter_edges_df = edges_df

                # ── GRAPH (expander, auto-opens when a row is selected) ────────
                graph_label = (
                    f"🔵 Subject network — {selected_snc_name}"
                    if selected_snc_name else "🔵 View subject network (all organisations)"
                )
                with st.expander(graph_label, expanded=selected_snc_name is not None):
                    html = build_similar_no_collab_graph(filter_similar_df, filter_edges_df)
                    html = inject_layout_controls(inject_png_download(html, filename="potential_partners.png"))
                    components.html(html, height=780, scrolling=True)
                    st.caption(
                        f"🔵 Blue nodes = potential partners  "
                        f"🟠 Orange nodes = shared research subjects.  "
                        "Hover over any node or edge for details."
                    )
                    graph_csv_df = filter_similar_df[[
                        "ORG_NAME", "ORG_CATEGORY",
                        "SHARED_SUBJECTS", "TOTAL_WEIGHT",
                        "SHARED_SUBJECT_NAMES",
                    ]].copy()
                    graph_csv_df.columns = [
                        "Organisation", "Category",
                        "Shared Subjects", "Patents & Publications",
                        "Shared Subject Areas",
                    ]
                    graph_csv_df.index = range(1, len(graph_csv_df) + 1)
                    st.download_button(
                        label="⬇️ Download CSV",
                        data=graph_csv_df.to_csv().encode("utf-8"),
                        file_name="potential_partners_network.csv",
                        mime="text/csv",
                        key="dl_snc_graph",
                    )

    else:
        # --- Standard graph mode ---
        st.markdown(
            "Nodes represent institutions, corporations, or research subjects. "
            "Thicker lines indicate stronger collaboration. **Hover over any node or edge** for details."
        )

        where_clauses = [f"WEIGHT >= {min_weight}"]

        if selected_ip_type != "All":
            where_clauses.append(f"IP_TYPE = '{sql_escape(selected_ip_type)}'")

        if subject_filter:
            safe_subj = sql_escape(subject_filter)
            where_clauses.append(f"(SOURCE_NAME ILIKE '%{safe_subj}%' OR TARGET_NAME ILIKE '%{safe_subj}%')")

        if selected_category != "All":
            safe_cat = sql_escape(selected_category)
            where_clauses.append(f"(SOURCE_CATEGORY = '{safe_cat}' OR TARGET_CATEGORY = '{safe_cat}')")

        if search_term.strip():
            safe_search = sql_escape(search_term.strip())
            where_clauses.append(f"(SOURCE_NAME ILIKE '%{safe_search}%' OR TARGET_NAME ILIKE '%{safe_search}%')")

        where_sql = " AND ".join(where_clauses)

        sql = f"""
SELECT
    SOURCE,
    SOURCE_NAME,
    SOURCE_TYPE,
    SOURCE_CATEGORY,
    SOURCE_NUS_AFFILIATED,
    TARGET,
    TARGET_NAME,
    TARGET_TYPE,
    TARGET_CATEGORY,
    TARGET_NUS_AFFILIATED,
    EDGE_TYPE,
    IP_TYPE,
    WEIGHT
FROM GRAPH_NETWORK.GRAPH.ALL_EDGES_ENRICHED
WHERE {where_sql}
ORDER BY WEIGHT DESC
LIMIT {max_edges}
"""

        df = run_query(sql)

        if top_n_nodes and search_term.strip():
            df = keep_top_n_neighbours(df, search_term.strip(), top_n_nodes)
        elif not search_term.strip():
            df = keep_top_communities(df, top_n=10)

        if df.empty:
            st.info(
                "No connections found for the selected filters. "
                "Try asking the AI assistant."
            )
        else:
            highlight_term = st.text_input(
                "🔍 Highlight a node",
                placeholder="Type a node name to highlight it in the graph…",
                key="highlight_input",
            )

            html = build_pyvis_graph(df, highlight_term=highlight_term.strip() if highlight_term else None)
            html = inject_layout_controls(inject_png_download(html, filename="collaboration_network.png"))
            components.html(html, height=780, scrolling=True)

            n_nodes = pd.concat([df["SOURCE"], df["TARGET"]]).nunique()
            n_edges = len(df)
            st.caption(
                f"Showing {n_nodes:,} institutions across {n_edges:,} connections. "
                "Drag nodes to rearrange, scroll to zoom."
            )

        with st.expander("📊 View connection data"):
            st.dataframe(df, use_container_width=True)
            col1, col2 = st.columns([1, 5])
            with col1:
                st.download_button(
                    label="⬇️ Download CSV",
                    data=df.to_csv(index=False).encode("utf-8"),
                    file_name="collaboration_network.csv",
                    mime="text/csv",
                )

        with st.expander("🔍 View SQL query"):
            st.code(sql, language="sql")

with chat_col:
    st.subheader("💬 AI Research Assistant")
    st.caption("Ask in plain language — or click an example to get started:")
    _CHAT_EXAMPLES = [
        "Recommend top 5 industry partners for NUS in AI",
    ]
    _ex_cols = st.columns(len(_CHAT_EXAMPLES))
    for _i, _ex in enumerate(_CHAT_EXAMPLES):
        with _ex_cols[_i]:
            if st.button(_ex, key=f"chat_ex_{_i}", use_container_width=True):
                st.session_state["pending_chat_q"] = _ex

    chat_container = st.container(height=400)
    with chat_container:
        for msg in st.session_state.chat_display:
            if msg["role"] == "user":
                st.chat_message("user").write(msg["content"])
            else:
                with st.chat_message("assistant"):
                    st.write(msg["content"])
                    if msg.get("filters"):
                        with st.expander("Applied filters", expanded=False):
                            st.json(msg["filters"])

    # Chat input using st.form — Enter submits, clears after send, renders inline
    with st.form(key="chat_form", clear_on_submit=True):
        col1, col2 = st.columns([5, 1])
        with col1:
            user_input = st.text_input(
                "Your question",
                placeholder="e.g. Recommend industry partners for NUS in AI",
                label_visibility="collapsed",
            )
        with col2:
            submitted = st.form_submit_button("Send ➤", use_container_width=True)

_pending_chat = st.session_state.pop("pending_chat_q", None)
if _pending_chat:
    user_input = _pending_chat
    submitted = True

if submitted and user_input.strip():
    st.session_state.chat_display.append({"role": "user", "content": user_input})

    with st.spinner("Thinking…"):
        try:
            parsed = extract_filters_from_llm(
                user_message=user_input,
                chat_history=st.session_state.chat_history,
                available_ip_types=ip_types_raw,
                available_edge_types=edge_types_raw,
                available_categories=categories_raw,
                available_subjects=subjects_raw,
            )

            response_type = parsed.get("response_type", "graph_query")

            if response_type == "general_answer":
                # --- General knowledge answer — don't touch filters ---
                answer = parsed.get("answer", "I'm not sure about that. Try rephrasing your question.")
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
                st.session_state.chat_display.append({
                    "role": "assistant",
                    "content": answer,
                    "filters": None,
                })
                log_query("collaborator_finder", user_input, "general_answer", llm_answer=answer)

            elif response_type == "recommendation":
                # --- Recommendation mode (flat model on PAT_PUB + ENTITIES) ---
                WRITE_UP_LIMIT = 10  # max orgs with detailed AI write-ups
                subject_val = parsed.get("subject_filter")
                existing_only = bool(parsed.get("existing_only", False))
                # existing_only queries expect more results by default (show the landscape)
                default_n = 15 if existing_only else 10
                top_n_val = int(parsed.get("top_n_results") or default_n)

                with st.spinner("Scoring partners on Snowflake…"):
                    recs_df = run_recommendation_flat(
                        subject=subject_val,
                        existing_only=existing_only,
                        top_n=top_n_val,
                    )

                if recs_df.empty:
                    answer = "No matching organisations found. Try a broader or differently-worded field."
                else:
                    # Write-ups capped at WRITE_UP_LIMIT; summary table shows all rows
                    write_up_df = recs_df.head(WRITE_UP_LIMIT)

                    with st.spinner("Fetching relevant titles…"):
                        titles_df = run_titles_for_flat_orgs(
                            tuple(write_up_df["ORG_NAME"].tolist()),
                            subject_val,
                        )

                    with st.spinner("Generating recommendations…"):
                        rec_text = generate_recommendations_flat(
                            write_up_df,
                            subject_val,
                            existing_only,
                            titles_df=titles_df,
                        )

                    if len(recs_df) > WRITE_UP_LIMIT:
                        note = (
                            f"_Detailed write-ups are shown for the **top {len(write_up_df)}** of "
                            f"**{len(recs_df)}** ranked partners. "
                            f"See the summary table on the right for the full ranked list._\n\n---\n\n"
                        )
                        rec_text = note + rec_text

                    # Store recommendation data for the flat-model summary view
                    st.session_state.filter_state["recommendation_data"] = {
                        "recs_df": recs_df.to_dict("records"),
                        "institution": "National University of Singapore",
                        "subject_filter": subject_val,
                        "rec_text": rec_text,
                        "existing_only": existing_only,
                        "flat": True,
                    }
                    st.session_state.filter_state["has_queried"] = True
                    st.session_state.filter_state["query_mode"] = "recommendation"

                    answer = rec_text

                st.session_state.chat_history.append({"role": "user", "content": user_input})
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
                st.session_state.chat_display.append({
                    "role": "assistant",
                    "content": answer,
                    "filters": None,
                })
                log_query(
                    "collaborator_finder", user_input, "recommendation",
                    subject_filter=subject_val, existing_only=existing_only, llm_answer=answer,
                    result_count=(0 if recs_df.empty else len(recs_df)),
                    result_orgs=(None if recs_df.empty
                                 else ", ".join(recs_df["ORG_NAME"].head(10).astype(str).tolist())),
                )

            else:
                # --- Graph query — update filters as normal ---
                explanation = parsed.pop("explanation", "Filters updated.")
                parsed.pop("answer", None)
                parsed.pop("response_type", None)

                new_fs = apply_llm_filters(
                    parsed,
                    st.session_state.filter_state,
                    ip_types_raw,
                    edge_types_raw,
                    categories_raw,
                )
                new_fs["has_queried"] = True
                st.session_state.filter_state = new_fs

                changed = {k: v for k, v in parsed.items() if v is not None}

                st.session_state.chat_history.append({"role": "user", "content": user_input})
                st.session_state.chat_history.append({"role": "assistant", "content": explanation})
                st.session_state.chat_display.append({
                    "role": "assistant",
                    "content": explanation,
                    "filters": changed if changed else None,
                })
                log_query("collaborator_finder", user_input, "graph_query",
                          subject_filter=changed.get("subject_filter"), llm_answer=explanation)

                run_query.clear()

        except Exception as e:
            error_msg = f"Sorry, I couldn't understand that request. Please try rephrasing. (Error: {e})"
            st.session_state.chat_display.append({"role": "assistant", "content": error_msg})

    st.rerun()
